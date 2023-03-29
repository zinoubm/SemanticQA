from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import Paragraph
import docx


def dummy_pdf(text: str):
    temp_file_path = "/tmp/temp_file"

    # Set up the PDF canvas and page size
    c = canvas.Canvas(temp_file_path, pagesize=letter)

    # Create a style for the paragraph
    style = ParagraphStyle(
        name="Normal", alignment=TA_CENTER, fontName="Helvetica", fontSize=12
    )

    # Add Lorem text to the PDF
    p = Paragraph(text, style=style)
    p.wrapOn(c, inch * 6, inch * 4)
    p.drawOn(c, inch * 2, inch * 5)

    # Save the PDF file
    c.save()

    return temp_file_path


def dummy_txt(text: str):
    temp_file_path = "/tmp/temp_file"

    with open(temp_file_path, "w") as file:
        file.write(text)

    return temp_file_path


def dummy_docx(text: str):
    # temp_file_path = "/tmp/temp_file"
    temp_file_path = "text.docx"

    # create a new Word document
    document = docx.Document()

    # add a heading to the document
    document.add_heading("My Document Title", level=0)

    # add a paragraph to the document
    paragraph = document.add_paragraph(text)

    # add a bold text run to the paragraph
    bold_text = paragraph.add_run(text)
    bold_text.bold = True

    # add a table to the document
    table = document.add_table(rows=3, cols=3)
    for i in range(3):
        row_cells = table.rows[i].cells
        for j in range(3):
            row_cells[j].text = f"Row {i+1}, Column {j+1}"

    # save the document to disk
    document.save(temp_file_path)

    return temp_file_path


if __name__ == "__main__":
    dummy_docx("lorem ipsum wow this is working!")
